# app.py
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy import stats
import math
import re
import base64
from io import BytesIO
from typing import List, Dict, Any, Tuple, Optional
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json
import os
import secrets
import string

st.set_page_config(page_title="Para_Variation - 蒙特卡洛模拟", layout="wide")

# ==================== 多语言文本字典 ====================
TEXTS = {
    "zh": {
        "title": "📊 Para_Variation - 基于蒙特卡洛模拟分析",
        "subtitle": "根据输入参数的分布进行随机抽样，计算用户定义的公式结果，分析输出分布及各参数贡献度。",
        "sim_settings": "⚙️ 模拟设置",
        "trail_number": "模拟次数 (Trail number)",
        "spec_limits": "规格限（可留空）",
        "usl": "规格上限 (USL)",
        "lsl": "规格下限 (LSL)",
        "random_seed": "随机种子",
        "about_system": "关于分析系统",
        "about_desc1": "**设计变量**：根据输入参数的公式计算得出。",
        "about_desc2": "**参数抽样**：每个参数独立根据其概率密度函数（PDF）进行随机抽样。",
        "output_title": "**输出结果：**",
        "output1": "- 预测设计变量在量产阶段的分布形态、均值及失效率。",
        "output2": "- 量化各设计参数对输出变量的影响百分比。",
        "output3": "- 通过调节规格上下限，快速确定合理的失效率目标，从而定义合理的工程规格。",
        "analyst_info": "分析人信息",
        "analyst_name": "分析人姓名",
        "analyst_title": "分析人头衔（可选）",
        "contact": "联系：",
        "email": "电邮: Techlife2027@gmail.com",
        "param_input": "📝 参数输入",
        "letter": "字母",
        "param_name": "参数名称",
        "mean": "均值(Typ)",
        "std": "标准差(Std)",
        "distribution": "分布",
        "delete": "删除",
        "add_row": "➕ 添加参数行",
        "new_param_default": "新参数",
        "configure": "⚙️ 配置 {} 参数",
        "dist_full": "正态分布（完整）",
        "dist_pos": "正态分布（正值）",
        "dist_neg": "正态分布（负值）",
        "dist_uniform": "均匀分布",
        "dist_lognorm": "对数正态分布",
        "dist_weibull": "威布尔分布",
        "dist_tri": "三角分布",
        "uniform_low": "下限",
        "uniform_high": "上限",
        "lognorm_meanlog": "对数均值 (μ_log)",
        "lognorm_sigmalog": "对数标准差 (σ_log)",
        "weibull_shape": "形状参数 (k)",
        "weibull_scale": "尺度参数 (λ)",
        "tri_left": "最小值",
        "tri_mode": "最可能值",
        "tri_right": "最大值",
        "error_low_high": "下限必须小于上限",
        "error_sigma": "对数标准差必须大于0",
        "error_weibull": "形状和尺度参数必须 > 0",
        "error_tri": "必须满足：最小值 ≤ 最可能值 ≤ 最大值",
        "formula_def": "📐 公式定义（设计值）",
        "design_var_name": "📌 设计变量名称",
        "formula_label": "📝 计算公式",
        "formula_hint": "💡 可直接在公式中使用字母（A, B, C...）代表对应参数，系统将自动识别。例如：A*E*7/1000*60/(B+C+D)",
        "formula_supported": "支持的运算: + - * / **, 括号, 函数: sqrt, exp, log, sin, cos, tan, pi, e 等。公式中的空格会被自动优化。",
        "design_value": "📌 当前设计值（基于均值）:",
        "formula_invalid": "公式无效或参数不匹配，无法计算设计值。请检查公式中的字母是否与上方对应关系一致，并确保运算正确。",
        "start_sim": "开始\n蒙特卡洛模拟",
        "sim_result": "📈 模拟结果: {}",
        "mean_val": "{} 均值",
        "std_val": "{} 标准差",
        "max_val": "最大值",
        "min_val": "最小值",
        "cpk_val": "Cpk",
        "failure_ppm": "失效率 - ppm level",
        "ppm_hint": "💡 可调节上下限以实时观察PPM水平的变化（留空表示无此限）",
        "no_limits": "未提供任何规格限，无法计算CPK和PPM。",
        "histogram": "{} 分布直方图",
        "hist_caption": "横轴：{}   |   纵轴：频次",
        "effect_chart": "{} 设计参数影响百分比",
        "effect_caption": "横轴：影响百分比   |   纵轴：设计参数",
        "view_contrib": "查看贡献百分比数据表",
        "view_data": "查看全部模拟数据",
        "download_csv": "📥 下载模拟数据 (CSV)",
        "download_report": "📄 下载专业报告 (Word)",
        "success": "模拟完成！",
        "report_title": "{} - DFSS模拟分析报告",
        "analyst_info_report": "分析人信息",
        "analyst_name_report": "分析人姓名：",
        "analyst_title_report": "头衔：",
        "not_filled": "未填写",
        "sim_settings_report": "1. 模拟设置",
        "output_var": "输出变量名称：",
        "formula_report": "公式：",
        "sim_times": "模拟次数：",
        "random_seed_report": "随机种子：",
        "usl_report": "规格上限 (USL)：",
        "lsl_report": "规格下限 (LSL)：",
        "none": "无",
        "param_table": "2. 输入参数表",
        "result_stats": "3. {}模拟结果统计",
        "statistic": "统计量",
        "value": "数值",
        "mean_stat": "均值",
        "std_stat": "标准差",
        "max_stat": "最大值",
        "min_stat": "最小值",
        "cpk_stat": "Cpk",
        "fail_all": "Failure All (ppm)",
        "fail_up": "Failure Up (ppm)",
        "fail_dn": "Failure Dn (ppm)",
        "histogram_report": "4. 分布直方图",
        "effect_report": "5. 设计参数对 {} 影响百分比",
        "detail_table": "详细数据表",
        "param": "参数",
        "contribution": "贡献百分比",
        "contact_report": "联系电邮：Techlife2027@gmail.com",
        "report_date": "报告生成时间：{}",
        "license_info": "授权信息",
        "remaining_label": "剩余次数",
        "expiry_label": "有效期至",
        "report_key_label": "授权码 (Report Key)",
        "no_license": "未输入授权码，当前为试用模式（本次会话剩余次数：{}）",
        "trial_warning": "⚠️ 您还有 {} 次试用机会，输入授权码可解锁无限使用和下载功能。",
        "purchase_button": "💰 购买授权码",
        "need_license": "⚠️ 请先购买授权码后再使用模拟功能。",
        "analyze_disabled": "您的免费次数已用完，请购买授权码后继续使用。",
        "purchase_dialog_title": "购买授权码",
        "plan_single": "单次通行",
        "plan_50": "50次套餐",
        "plan_1000": "1000次套餐",
        "payment_note": "支付成功后，您将收到授权码。请将授权码粘贴到左侧边栏输入框中即可解锁全部功能。",
        "payment_success_title": "✅ 支付成功",
        "payment_success_msg": "您的授权码已生成",
        "payment_save_key": "请妥善保管此授权码，下次使用时可手动复制并粘贴到左侧输入框。",
        "admin_settings": "管理员设置",
        "admin_login": "管理员验证",
        "username": "用户名",
        "password": "密码",
        "login": "登录",
        "key_generator": "Report Key 生成器",
        "license_type": "选择授权类型",
        "custom_uses": "使用次数",
        "custom_months": "有效期（月）",
        "custom_key_input": "自定义授权码（可选，留空则自动生成）",
        "generate_key": "生成 Report Key",
        "generated_key": "已生成 {} Report Key：",
        "key_list": "已生成的所有 Report Key",
        "show_limit": "显示条数",
        "export_keys": "📥 导出所有授权码为 Excel",
        "no_keys": "暂无授权码记录",
        "close": "确定",
    },
    "en": {
        "title": "📊 Para_Variation - Monte Carlo Simulation",
        "subtitle": "Randomly sample input parameters based on their distributions, compute user-defined formula, analyze output distribution and parameter contribution.",
        "sim_settings": "⚙️ Simulation Settings",
        "trail_number": "Trail number",
        "spec_limits": "Specification limits (leave blank if none)",
        "usl": "Upper Spec Limit (USL)",
        "lsl": "Lower Spec Limit (LSL)",
        "random_seed": "Random seed",
        "about_system": "About Analysis System",
        "about_desc1": "**Design Variable**: Calculated from input parameters using the formula.",
        "about_desc2": "**Parameter Sampling**: Each parameter is independently sampled based on its Probability Density Function (PDF).",
        "output_title": "**Outputs:**",
        "output1": "- Predict the distribution shape, mean, and failure rate of the design variable in mass production.",
        "output2": "- Quantify the percentage impact of each design parameter on the output variable.",
        "output3": "- Adjust spec limits to quickly determine reasonable failure rate targets and define engineering specifications.",
        "analyst_info": "Analyst Information",
        "analyst_name": "Analyst Name",
        "analyst_title": "Analyst Title (optional)",
        "contact": "Contact:",
        "email": "Email: Techlife2027@gmail.com",
        "param_input": "📝 Parameter Input",
        "letter": "Letter",
        "param_name": "Parameter Name",
        "mean": "Mean (Typ)",
        "std": "Std Dev (Std)",
        "distribution": "Distribution",
        "delete": "Delete",
        "add_row": "➕ Add Parameter Row",
        "new_param_default": "New Parameter",
        "configure": "⚙️ Configure {} Parameters",
        "dist_full": "Normal (Full)",
        "dist_pos": "Normal (Positive only)",
        "dist_neg": "Normal (Negative only)",
        "dist_uniform": "Uniform",
        "dist_lognorm": "Log-normal",
        "dist_weibull": "Weibull",
        "dist_tri": "Triangular",
        "uniform_low": "Lower bound",
        "uniform_high": "Upper bound",
        "lognorm_meanlog": "Log mean (μ_log)",
        "lognorm_sigmalog": "Log std dev (σ_log)",
        "weibull_shape": "Shape (k)",
        "weibull_scale": "Scale (λ)",
        "tri_left": "Minimum",
        "tri_mode": "Most likely",
        "tri_right": "Maximum",
        "error_low_high": "Lower bound must be less than upper bound",
        "error_sigma": "Log standard deviation must be > 0",
        "error_weibull": "Shape and scale must be > 0",
        "error_tri": "Must satisfy: min ≤ mode ≤ max",
        "formula_def": "📐 Formula Definition (Design Value)",
        "design_var_name": "📌 Design Variable Name",
        "formula_label": "📝 Formula",
        "formula_hint": "💡 Use letters (A, B, C...) in the formula to represent parameters. System will automatically recognize. Example: A*E*7/1000*60/(B+C+D)",
        "formula_supported": "Supported operators: + - * / **, parentheses, functions: sqrt, exp, log, sin, cos, tan, pi, e, etc. Spaces are automatically optimized.",
        "design_value": "📌 Current Design Value (based on means):",
        "formula_invalid": "Invalid formula or parameter mismatch. Cannot compute design value. Please check if letters match the table above.",
        "start_sim": "Start\nMonte Carlo",
        "sim_result": "📈 Simulation Results: {}",
        "mean_val": "{} Mean",
        "std_val": "{} Std Dev",
        "max_val": "Max",
        "min_val": "Min",
        "cpk_val": "Cpk",
        "failure_ppm": "Failure ppm level",
        "ppm_hint": "💡 Adjust limits to see PPM levels in real time (leave blank if no limit).",
        "no_limits": "No specification limits provided. Cannot compute CPK and PPM.",
        "histogram": "{} Distribution Histogram",
        "hist_caption": "X-axis: {}   |   Y-axis: Frequency",
        "effect_chart": "Parameter Contribution to {}",
        "effect_caption": "X-axis: Contribution %   |   Y-axis: Parameters",
        "view_contrib": "View Contribution Data Table",
        "view_data": "View All Simulation Data",
        "download_csv": "📥 Download Simulation Data (CSV)",
        "download_report": "📄 Download Professional Report (Word)",
        "success": "Simulation completed!",
        "report_title": "{} - DFSS Simulation Report",
        "analyst_info_report": "Analyst Information",
        "analyst_name_report": "Analyst Name:",
        "analyst_title_report": "Title:",
        "not_filled": "Not filled",
        "sim_settings_report": "1. Simulation Settings",
        "output_var": "Output Variable Name:",
        "formula_report": "Formula:",
        "sim_times": "Number of simulations:",
        "random_seed_report": "Random seed:",
        "usl_report": "Upper Spec Limit (USL):",
        "lsl_report": "Lower Spec Limit (LSL):",
        "none": "None",
        "param_table": "2. Input Parameter Table",
        "result_stats": "3. {} Simulation Result Statistics",
        "statistic": "Statistic",
        "value": "Value",
        "mean_stat": "Mean",
        "std_stat": "Std Dev",
        "max_stat": "Max",
        "min_stat": "Min",
        "cpk_stat": "Cpk",
        "fail_all": "Failure All (ppm)",
        "fail_up": "Failure Up (ppm)",
        "fail_dn": "Failure Dn (ppm)",
        "histogram_report": "4. Distribution Histogram",
        "effect_report": "5. Parameter Contribution to {}",
        "detail_table": "Detail Data Table",
        "param": "Parameter",
        "contribution": "Contribution %",
        "contact_report": "Contact Email: Techlife2027@gmail.com",
        "report_date": "Report Date: {}",
        "license_info": "License Info",
        "remaining_label": "Remaining uses",
        "expiry_label": "Valid until",
        "report_key_label": "Report Key",
        "no_license": "No Report Key. Trial mode (remaining credits this session: {})",
        "trial_warning": "⚠️ You have {} trial credits left. Enter a license key to unlock unlimited usage.",
        "purchase_button": "💰 Purchase License",
        "need_license": "⚠️ Please purchase a license before using simulation.",
        "analyze_disabled": "Your free trial has expired. Please purchase a license to continue.",
        "purchase_dialog_title": "Purchase License",
        "plan_single": "Single Pass",
        "plan_50": "50 Credits",
        "plan_1000": "1000 Credits",
        "payment_note": "After successful payment, you will receive a license key. Please paste it into the left sidebar to unlock all features.",
        "payment_success_title": "✅ Payment Successful",
        "payment_success_msg": "Your license key has been generated",
        "payment_save_key": "Please save this license key. You can copy and paste it into the left sidebar next time.",
        "admin_settings": "Admin Settings",
        "admin_login": "Admin Verification",
        "username": "Username",
        "password": "Password",
        "login": "Login",
        "key_generator": "Report Key Generator",
        "license_type": "License Type",
        "custom_uses": "Number of uses",
        "custom_months": "Validity (months)",
        "custom_key_input": "Custom license key (optional, leave blank to auto-generate)",
        "generate_key": "Generate Report Key",
        "generated_key": "Generated {} Report Key:",
        "key_list": "All Generated Report Keys",
        "show_limit": "Show",
        "export_keys": "📥 Export all keys to Excel",
        "no_keys": "No license keys yet.",
        "close": "OK",
    }
}

# ==================== 初始化 Session State ====================
if "lang" not in st.session_state:
    st.session_state.lang = "zh"
if "admin_logged_in" not in st.session_state:
    st.session_state.admin_logged_in = False
if "analyst_name" not in st.session_state:
    st.session_state.analyst_name = ""
if "analyst_title" not in st.session_state:
    st.session_state.analyst_title = ""
if "current_report_key" not in st.session_state:
    st.session_state.current_report_key = ""
if "trial_uses_left" not in st.session_state:
    st.session_state.trial_uses_left = 3
if "sim_results_raw" not in st.session_state:
    st.session_state.sim_results_raw = None
if "formula" not in st.session_state:
    st.session_state.formula = "A * E * 7 / 1000 * 60 / (B + C + D)"
if "output_name" not in st.session_state:
    st.session_state.output_name = "Runtime"
if "usl_str" not in st.session_state:
    st.session_state.usl_str = "40.0"
if "lsl_str" not in st.session_state:
    st.session_state.lsl_str = "30.0"
if "params" not in st.session_state:
    st.session_state.params = pd.DataFrame({
        "参数名称": ["Cell Cap", "Suction P", "Brush P", "Other(Pump+display)", "V"],
        "均值(Typ)": [2450.0, 70.0, 30.0, 15.0, 3.6],
        "标准差(Std)": [20.74, 0.77, 0.90, 0.45, 0.0036],
        "分布": ["正态分布（完整）" for _ in range(5)],
        "分布参数": [{} for _ in range(5)]
    })
if "show_payment_dialog" not in st.session_state:
    st.session_state.show_payment_dialog = False
if "payment_new_key" not in st.session_state:
    st.session_state.payment_new_key = ""

# ==================== 辅助函数 ====================
def t(key):
    return TEXTS[st.session_state.lang].get(key, key)

# ==================== 授权与试用数据管理 ====================
USAGE_FILE = "usage_data.json"

def load_usage_data():
    if os.path.exists(USAGE_FILE):
        try:
            with open(USAGE_FILE, "r") as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_usage_data(data):
    with open(USAGE_FILE, "w") as f:
        json.dump(data, f, indent=2)

LICENSE_TYPES = {
    "trial": {"name": "试用版", "max_uses": 3, "max_months": 1, "en_name": "Trial"},
    "level1": {"name": "一级用户", "max_uses": 100, "max_months": 12, "en_name": "Level 1"},
    "level2": {"name": "二级用户", "max_uses": 300, "max_months": 24, "en_name": "Level 2"},
    "level3": {"name": "三级用户", "max_uses": 500, "max_months": 36, "en_name": "Level 3"},
    "level4": {"name": "四级用户", "max_uses": 1000, "max_months": 60, "en_name": "Level 4"},
}

def generate_report_key(license_type, custom_uses=None, custom_months=None, custom_key=None):
    if license_type == "custom":
        max_uses = custom_uses
        max_months = custom_months
        type_name = "自定义"
    else:
        lic_info = LICENSE_TYPES[license_type]
        max_uses = lic_info["max_uses"]
        max_months = lic_info["max_months"]
        type_name = lic_info["name"]
    expiry = datetime.now() + timedelta(days=max_months*30)
    expiry_str = expiry.isoformat()
    usage_db = load_usage_data()
    if custom_key and custom_key.strip():
        new_key = custom_key.strip().upper()
        if new_key in usage_db:
            return None, 0, None, "授权码已存在"
    else:
        while True:
            random_str = ''.join(secrets.choice(string.ascii_uppercase + string.digits) for _ in range(8))
            new_key = f"{license_type.upper()}_{random_str}"
            if new_key not in usage_db:
                break
    usage_db[new_key] = {
        "type": license_type,
        "remaining": max_uses,
        "expiry": expiry_str,
        "total_uses": 0,
        "generated_at": datetime.now().isoformat()
    }
    save_usage_data(usage_db)
    return new_key, max_uses, expiry_str, type_name

def activate_license(report_key):
    if not report_key:
        return False, 0, None, None
    usage_db = load_usage_data()
    if report_key in usage_db:
        record = usage_db[report_key]
        remaining = record["remaining"]
        expiry_str = record["expiry"]
        expiry = datetime.fromisoformat(expiry_str)
        if remaining > 0 and datetime.now() <= expiry:
            return True, remaining, expiry_str, record.get("type", "unknown")
    return False, 0, None, None

def consume_usage(report_key):
    if st.session_state.get("admin_logged_in", False):
        return True
    if not report_key:
        if st.session_state.trial_uses_left > 0:
            st.session_state.trial_uses_left -= 1
            return True
        else:
            return False
    usage_db = load_usage_data()
    if report_key in usage_db:
        record = usage_db[report_key]
        if record["remaining"] > 0 and datetime.now() <= datetime.fromisoformat(record["expiry"]):
            record["remaining"] -= 1
            record["total_uses"] = record.get("total_uses", 0) + 1
            save_usage_data(usage_db)
            return True
    return False

def get_remaining_info(report_key):
    if st.session_state.get("admin_logged_in", False):
        return ("无限" if st.session_state.lang=="zh" else "Unlimited"), ("永久" if st.session_state.lang=="zh" else "Permanent")
    if report_key:
        valid, remaining, expiry_str, _ = activate_license(report_key)
        if valid:
            return str(remaining), expiry_str[:10]
    return str(st.session_state.trial_uses_left), ("试用剩余次数" if st.session_state.lang=="zh" else "Trial left")

def is_premium_user(report_key):
    if st.session_state.get("admin_logged_in", False):
        return True
    if report_key:
        valid, _, _, _ = activate_license(report_key)
        return valid
    return False

# ==================== 支付链接配置 ====================
PAYMENT_LINKS = {
    "single": {
        "url": "https://buy.stripe.com/test_7sY8wPcYJ5Qu898cjO6Vq00",
        "name_zh": "单次通行",
        "name_en": "Single Pass",
        "price_usd": 3,
        "uses": 3,
        "months": 9999
    },
    "50": {
        "url": "https://buy.stripe.com/test_cNi3cv1g1a6KfBA6Zu6Vq01",
        "name_zh": "50次套餐",
        "name_en": "50 Credits",
        "price_usd": 30,
        "uses": 50,
        "months": 1
    },
    "1000": {
        "url": "https://buy.stripe.com/test_00wfZh6Alen0ahg2Je6Vq02",
        "name_zh": "1000次套餐",
        "name_en": "1000 Credits",
        "price_usd": 200,
        "uses": 1000,
        "months": 12
    }
}

# ==================== 管理员登录验证 ====================
ADMIN_USERNAME = "Laurence_ku"
ADMIN_PASSWORD = "Ku_product$2026"

# ==================== 支付回调处理 ====================
def handle_payment_callback():
    params = st.query_params
    if "order_success" in params and "plan" in params:
        plan_key = params["plan"]
        if plan_key in PAYMENT_LINKS:
            uses = PAYMENT_LINKS[plan_key]["uses"]
            months = PAYMENT_LINKS[plan_key]["months"]
            new_key, max_uses, expiry_str, _ = generate_report_key("custom", custom_uses=uses, custom_months=months)
            if new_key:
                st.session_state.current_report_key = new_key
                st.session_state.payment_new_key = new_key
                st.session_state.show_payment_dialog = True
                st.query_params.clear()
                st.rerun()
            else:
                st.error("生成授权码失败，请联系管理员。" if st.session_state.lang=="zh" else "Failed to generate license key. Contact admin.")
                st.query_params.clear()
        else:
            st.error("无效的套餐类型。" if st.session_state.lang=="zh" else "Invalid plan type.")
            st.query_params.clear()

def show_payment_success_dialog():
    if st.session_state.get("show_payment_dialog", False):
        @st.dialog(t("payment_success_title") if st.session_state.lang=="zh" else "✅ Payment Successful")
        def payment_success_dialog():
            lang = st.session_state.lang
            st.markdown(f"### {t('payment_success_msg')}")
            st.code(st.session_state.payment_new_key, language="text")
            st.caption(t("payment_save_key"))
            if st.button(t("close")):
                st.session_state.show_payment_dialog = False
                st.session_state.payment_new_key = ""
                st.rerun()
        payment_success_dialog()

# ==================== 购买对话框 ====================
@st.dialog("购买授权码", width="large")
def purchase_dialog():
    lang = st.session_state.lang
    if lang == "zh":
        st.markdown("### 选择套餐")
        st.markdown("""
| 套餐 | 价格 | 次数 | 有效期 |
|------|------|------|--------|
| 单次通行 | 18元 / 3美元 | 3次 | 无限制 |
| 50次套餐 | 180元 / 30美元 | 50次 | 1个月 |
| 1000次套餐 | 1200元 / 200美元 | 1000次 | 12个月 |
""")
    else:
        st.markdown("### Select Plan")
        st.markdown("""
| Plan | Price | Credits | Validity |
|------|-------|---------|----------|
| Single Pass | 18 RMB / $3 | 3 uses | Unlimited |
| 50 Credits | 180 RMB / $30 | 50 uses | 1 month |
| 1000 Credits | 1200 RMB / $200 | 1000 uses | 12 months |
""")
    st.markdown("#### 💳 " + ("银行卡/数字钱包支付（Stripe）" if lang=="zh" else "Card / Digital Wallet Payment (Stripe)"))
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        url = PAYMENT_LINKS["single"]["url"]
        name = PAYMENT_LINKS["single"]["name_zh"] if lang=="zh" else PAYMENT_LINKS["single"]["name_en"]
        price = PAYMENT_LINKS["single"]["price_usd"]
        button_html = f'<a href="{url}" target="_blank" style="display: block; background-color: #E60000; color: white; font-weight: bold; font-size: 18px; padding: 12px; border-radius: 8px; text-align: center; text-decoration: none; width: 100%;">🎟️ {name} ${price}</a>'
        st.markdown(button_html, unsafe_allow_html=True)
    
    with col2:
        url = PAYMENT_LINKS["50"]["url"]
        name = PAYMENT_LINKS["50"]["name_zh"] if lang=="zh" else PAYMENT_LINKS["50"]["name_en"]
        price = PAYMENT_LINKS["50"]["price_usd"]
        button_html = f'<a href="{url}" target="_blank" style="display: block; background-color: #E60000; color: white; font-weight: bold; font-size: 18px; padding: 12px; border-radius: 8px; text-align: center; text-decoration: none; width: 100%;">📦 {name} ${price}</a>'
        st.markdown(button_html, unsafe_allow_html=True)
    
    with col3:
        url = PAYMENT_LINKS["1000"]["url"]
        name = PAYMENT_LINKS["1000"]["name_zh"] if lang=="zh" else PAYMENT_LINKS["1000"]["name_en"]
        price = PAYMENT_LINKS["1000"]["price_usd"]
        button_html = f'<a href="{url}" target="_blank" style="display: block; background-color: #E60000; color: white; font-weight: bold; font-size: 18px; padding: 12px; border-radius: 8px; text-align: center; text-decoration: none; width: 100%;">🚀 {name} ${price}</a>'
        st.markdown(button_html, unsafe_allow_html=True)
    
    st.markdown(t("payment_note"))

# ==================== 管理员设置弹窗 ====================
@st.dialog(t("admin_settings"), width="large")
def admin_settings_dialog():
    lang = st.session_state.lang
    st.subheader(t("admin_login"))
    if not st.session_state.admin_logged_in:
        username = st.text_input(t("username"))
        password = st.text_input(t("password"), type="password")
        if st.button(t("login")):
            if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
                st.session_state.admin_logged_in = True
                st.rerun()
            else:
                st.error("用户名或密码错误" if lang=="zh" else "Incorrect username or password")
        return

    st.success("管理员已登录" if lang=="zh" else "Admin logged in")
    st.markdown("---")
    st.subheader(t("key_generator"))
    key_type = st.selectbox(t("license_type"), ["试用版", "一级用户", "二级用户", "三级用户", "四级用户", "自定义"])
    custom_uses = None
    custom_months = None
    if key_type == "自定义":
        col_c1, col_c2 = st.columns(2)
        with col_c1:
            custom_uses = st.number_input(t("custom_uses"), min_value=1, step=1, value=100)
        with col_c2:
            custom_months = st.number_input(t("custom_months"), min_value=1, step=1, value=12)
    custom_key_input = st.text_input(t("custom_key_input"), placeholder="例如：VIP_2026_001")
    if st.button(t("generate_key")):
        if key_type == "试用版":
            lic_type = "trial"
        elif key_type == "一级用户":
            lic_type = "level1"
        elif key_type == "二级用户":
            lic_type = "level2"
        elif key_type == "三级用户":
            lic_type = "level3"
        elif key_type == "四级用户":
            lic_type = "level4"
        else:
            lic_type = "custom"
        result = generate_report_key(lic_type, custom_uses, custom_months, custom_key_input)
        if result[0] is None:
            st.error(result[3])
        else:
            new_key, max_uses, expiry_str, type_name = result
            st.success(t("generated_key").format(type_name))
            st.code(new_key, language="text")
            st.write(f"可使用次数：{max_uses} 次，有效期至：{expiry_str[:10]}" if lang=="zh" else f"Uses: {max_uses}, Valid until: {expiry_str[:10]}")
    
    st.markdown("---")
    st.subheader(t("key_list"))
    usage_db = load_usage_data()
    records = []
    for key, data in usage_db.items():
        gen_time = data.get("generated_at")
        if gen_time:
            try:
                gen_dt = datetime.fromisoformat(gen_time)
            except:
                gen_dt = datetime.min
        else:
            gen_dt = datetime.min
        records.append({
            "授权码" if lang=="zh" else "License Key": key,
            "类型" if lang=="zh" else "Type": data.get("type", "unknown"),
            "剩余次数" if lang=="zh" else "Remaining": data["remaining"],
            "总使用次数" if lang=="zh" else "Total uses": data.get("total_uses", 0),
            "有效期至" if lang=="zh" else "Valid until": data["expiry"][:10] if data["expiry"] else "永久",
            "生成时间" if lang=="zh" else "Generated at": gen_dt.strftime("%Y-%m-%d %H:%M:%S") if gen_dt != datetime.min else "未知"
        })
    records.sort(key=lambda x: x["生成时间" if lang=="zh" else "Generated at"], reverse=True)
    show_limit = st.selectbox(t("show_limit"), ["最近10条", "最近20条", "最近50条", "全部"] if lang=="zh" else ["Last 10", "Last 20", "Last 50", "All"], index=0)
    if lang=="zh":
        limit_map = {"最近10条":10, "最近20条":20, "最近50条":50, "全部":len(records)}
    else:
        limit_map = {"Last 10":10, "Last 20":20, "Last 50":50, "All":len(records)}
    limit = limit_map[show_limit]
    display_records = records[:limit]
    if display_records:
        df = pd.DataFrame(display_records)
        st.dataframe(df, use_container_width=True)
    else:
        st.info(t("no_keys"))
    if st.button(t("export_keys")):
        if records:
            df_all = pd.DataFrame(records)
            output = BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df_all.to_excel(writer, sheet_name="授权码列表" if lang=="zh" else "License Keys", index=False)
            excel_data = output.getvalue()
            st.download_button(label="点击下载 Excel 文件" if lang=="zh" else "Download Excel", data=excel_data, file_name=f"report_keys_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.warning(t("no_keys"))

# ==================== 蒙特卡洛模拟核心函数 ====================
def update_param_letters():
    letters = [chr(ord('A') + i) for i in range(len(st.session_state.params))]
    st.session_state.param_letters = {
        row["参数名称"]: letters[i] for i, row in st.session_state.params.iterrows()
    }
update_param_letters()

def parse_limit(s: str) -> Optional[float]:
    if s is None or s.strip() == "":
        return None
    try:
        return float(s)
    except ValueError:
        return None

def sync_usl_from_main(): st.session_state.usl_str = st.session_state.main_usl
def sync_lsl_from_main(): st.session_state.lsl_str = st.session_state.main_lsl
def sync_usl_from_sidebar(): st.session_state.usl_str = st.session_state.usl_sidebar
def sync_lsl_from_sidebar(): st.session_state.lsl_str = st.session_state.lsl_sidebar

def clean_formula(formula: str) -> str:
    formula = formula.strip()
    formula = re.sub(r'\s+', ' ', formula)
    formula = re.sub(r'(?<=[0-9a-zA-Z)])\s*([+\-*/])\s*(?=[0-9a-zA-Z(])', r' \1 ', formula)
    formula = re.sub(r'\(\s+', '(', formula)
    formula = re.sub(r'\s+\)', ')', formula)
    return formula

def replace_letters_with_names(expr: str, param_letters: Dict[str, str]) -> str:
    reverse_map = {v: k for k, v in param_letters.items()}
    for letter, name in sorted(reverse_map.items(), key=lambda x: len(x[0]), reverse=True):
        pattern = r'(?<![a-zA-Z0-9_])' + re.escape(letter) + r'(?![a-zA-Z0-9_])'
        expr = re.sub(pattern, name, expr)
    return expr

def safe_eval_with_mapping(expr: str, param_names: List[str], context_values: List[float], param_letters: Dict[str, str]) -> float:
    expr = clean_formula(expr)
    expr_with_names = replace_letters_with_names(expr, param_letters)
    temp_names = [f"__p{i}__" for i in range(len(param_names))]
    sorted_params = sorted(zip(param_names, temp_names), key=lambda x: len(x[0]), reverse=True)
    expr_temp = expr_with_names
    for orig, temp in sorted_params:
        pattern = r'(?<![a-zA-Z0-9_])' + re.escape(orig) + r'(?![a-zA-Z0-9_])'
        expr_temp = re.sub(pattern, temp, expr_temp)
    context = {temp: val for temp, val in zip(temp_names, context_values)}
    allowed_names = {
        "sqrt": math.sqrt, "exp": math.exp, "log": math.log, "log10": math.log10,
        "sin": math.sin, "cos": math.cos, "tan": math.tan, "pi": math.pi, "e": math.e,
        "abs": abs, "pow": pow
    }
    allowed_names.update(context)
    try:
        result = eval(expr_temp, {"__builtins__": {}}, allowed_names)
        return float(result)
    except Exception:
        return np.nan

def compute_design_value(params_df: pd.DataFrame, formula: str, param_letters: Dict[str, str]) -> Optional[float]:
    param_names = params_df["参数名称"].astype(str).tolist()
    means = params_df["均值(Typ)"].values.astype(float)
    val = safe_eval_with_mapping(formula, param_names, means, param_letters)
    return val if not np.isnan(val) else None

def get_distributions():
    return [
        t("dist_full"),
        t("dist_pos"),
        t("dist_neg"),
        t("dist_uniform"),
        t("dist_lognorm"),
        t("dist_weibull"),
        t("dist_tri")
    ]

def generate_sample(dist: str, mean: float, std: float, dist_params: Dict, size: int = 1) -> np.ndarray:
    if dist == t("dist_full"):
        return np.random.normal(mean, std, size)
    elif dist == t("dist_pos"):
        a, b = (0 - mean) / std if std > 0 else -np.inf, np.inf
        if std == 0:
            return np.full(size, max(mean, 0))
        return stats.truncnorm.rvs(a, b, loc=mean, scale=std, size=size)
    elif dist == t("dist_neg"):
        a, b = -np.inf, (0 - mean) / std if std > 0 else np.inf
        if std == 0:
            return np.full(size, min(mean, 0))
        return stats.truncnorm.rvs(a, b, loc=mean, scale=std, size=size)
    elif dist == t("dist_uniform"):
        low = dist_params.get("low", mean - 3*std)
        high = dist_params.get("high", mean + 3*std)
        return np.random.uniform(low, high, size)
    elif dist == t("dist_lognorm"):
        mean_log = dist_params.get("mean_log", 0.0)
        sigma_log = dist_params.get("sigma_log", 1.0)
        return np.random.lognormal(mean_log, sigma_log, size)
    elif dist == t("dist_weibull"):
        shape = dist_params.get("shape", 1.0)
        scale = dist_params.get("scale", 1.0)
        return np.random.weibull(shape, size) * scale
    elif dist == t("dist_tri"):
        left = dist_params.get("left", mean - 3*std)
        mode = dist_params.get("mode", mean)
        right = dist_params.get("right", mean + 3*std)
        return np.random.triangular(left, mode, right, size)
    else:
        return np.random.normal(mean, std, size)

def run_monte_carlo(params_df: pd.DataFrame, formula: str, n_sim: int, param_letters: Dict[str, str], seed: int = 38) -> Dict[str, Any]:
    np.random.seed(seed)
    n_params = len(params_df)
    param_names = params_df["参数名称"].astype(str).tolist()
    means = params_df["均值(Typ)"].values.astype(float)
    stds = params_df["标准差(Std)"].values.astype(float)
    dists = params_df["分布"].tolist()
    dist_params_list = params_df["分布参数"].tolist()

    samples = np.zeros((n_sim, n_params))
    for i in range(n_params):
        samples[:, i] = generate_sample(dists[i], means[i], stds[i], dist_params_list[i], n_sim)

    results = []
    for i in range(n_sim):
        val = safe_eval_with_mapping(formula, param_names, samples[i, :], param_letters)
        if not np.isnan(val):
            results.append(val)

    results = np.array(results)
    if len(results) == 0:
        st.error("所有公式计算均失败，请检查公式！")
        return None

    mean_out = np.mean(results)
    std_out = np.std(results, ddof=1)
    max_out = np.max(results)
    min_out = np.min(results)

    hist_counts, bin_edges = np.histogram(results, bins=25, density=False)
    bin_centers = (bin_edges[:-1] + bin_edges[1:]) / 2
    x_pdf = np.linspace(min_out, max_out, 200)
    pdf_theory = stats.norm.pdf(x_pdf, mean_out, std_out)

    return {
        "results": results,
        "samples": samples,
        "mean": mean_out,
        "std": std_out,
        "max": max_out,
        "min": min_out,
        "hist_counts": hist_counts,
        "bin_edges": bin_edges,
        "bin_centers": bin_centers,
        "x_pdf": x_pdf,
        "pdf_theory": pdf_theory,
        "param_names": param_names,
    }

def sensitivity_analysis(params_df: pd.DataFrame, formula: str, n_sim: int, param_letters: Dict[str, str], seed: int = 38) -> Tuple[pd.DataFrame, List[float], List[str]]:
    np.random.seed(seed)
    n_params = len(params_df)
    param_names = params_df["参数名称"].astype(str).tolist()
    means = params_df["均值(Typ)"].values.astype(float)
    stds = params_df["标准差(Std)"].values.astype(float)
    dists = params_df["分布"].tolist()
    dist_params_list = params_df["分布参数"].tolist()

    variances = []
    for i in range(n_params):
        samples_i = generate_sample(dists[i], means[i], stds[i], dist_params_list[i], n_sim)
        results_i = []
        for val in samples_i:
            context_vals = means.copy()
            context_vals[i] = val
            res = safe_eval_with_mapping(formula, param_names, context_vals, param_letters)
            if not np.isnan(res):
                results_i.append(res)
        var_i = np.var(results_i, ddof=1) if len(results_i) > 1 else 0.0
        variances.append(var_i)

    total_var = sum(variances)
    contributions = [v / total_var if total_var > 0 else 0.0 for v in variances]

    df_contrib = pd.DataFrame({
        "参数": param_names,
        "方差贡献": variances,
        "贡献百分比": contributions
    })
    df_contrib = df_contrib.sort_values("贡献百分比", ascending=False).reset_index(drop=True)
    df_contrib["贡献百分比_显示"] = df_contrib["贡献百分比"].apply(lambda x: f"{x:.1%}")
    return df_contrib, contributions, param_names

def plot_pdf(dist: str, mean: float, std: float, dist_params: Dict, ax):
    if dist == t("dist_full"):
        x = np.linspace(mean - 4*std, mean + 4*std, 200)
        y = stats.norm.pdf(x, mean, std)
        ax.plot(x, y, 'b-')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"N(μ={mean:.1f}, σ={std:.2f})", fontsize=8)
    elif dist == t("dist_pos"):
        a, b = (0 - mean) / std if std > 0 else -np.inf, np.inf
        if std == 0:
            x = [max(mean, 0)]
            y = [1]
        else:
            x = np.linspace(0, mean + 4*std, 200)
            y = stats.truncnorm.pdf(x, a, b, loc=mean, scale=std)
        ax.plot(x, y, 'g-')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"TruncNorm(≥0)", fontsize=8)
    elif dist == t("dist_neg"):
        a, b = -np.inf, (0 - mean) / std if std > 0 else np.inf
        if std == 0:
            x = [min(mean, 0)]
            y = [1]
        else:
            x = np.linspace(mean - 4*std, 0, 200)
            y = stats.truncnorm.pdf(x, a, b, loc=mean, scale=std)
        ax.plot(x, y, 'r-')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"TruncNorm(≤0)", fontsize=8)
    elif dist == t("dist_uniform"):
        low = dist_params.get("low", mean - 3*std)
        high = dist_params.get("high", mean + 3*std)
        x = np.linspace(low, high, 200)
        y = stats.uniform.pdf(x, low, high-low)
        ax.plot(x, y, 'purple')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"U({low:.1f}, {high:.1f})", fontsize=8)
    elif dist == t("dist_lognorm"):
        mean_log = dist_params.get("mean_log", 0.0)
        sigma_log = dist_params.get("sigma_log", 1.0)
        x = np.linspace(0, np.exp(mean_log + 3*sigma_log), 200)
        y = stats.lognorm.pdf(x, sigma_log, scale=np.exp(mean_log))
        ax.plot(x, y, 'orange')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"LogN(μlog={mean_log:.1f}, σlog={sigma_log:.2f})", fontsize=8)
    elif dist == t("dist_weibull"):
        shape = dist_params.get("shape", 1.0)
        scale = dist_params.get("scale", 1.0)
        x = np.linspace(0, scale * 3, 200)
        y = stats.weibull_min.pdf(x, shape, scale=scale)
        ax.plot(x, y, 'brown')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"Weibull(k={shape:.1f}, λ={scale:.1f})", fontsize=8)
    elif dist == t("dist_tri"):
        left = dist_params.get("left", mean - 3*std)
        mode = dist_params.get("mode", mean)
        right = dist_params.get("right", mean + 3*std)
        x = np.linspace(left, right, 200)
        y = stats.triang.pdf(x, (mode-left)/(right-left), loc=left, scale=right-left)
        ax.plot(x, y, 'olive')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"Tri({left:.1f}, {mode:.1f}, {right:.1f})", fontsize=8)
    ax.set_xlabel("Value", fontsize=6)
    ax.set_ylabel("Density", fontsize=6)
    ax.tick_params(axis='both', labelsize=6)

def plot_histogram(results, bin_centers, hist_counts, x_pdf, pdf_theory, usl, lsl, output_name, n_sim):
    fig, ax = plt.subplots(figsize=(11, 6), dpi=100)
    ax.bar(bin_centers, hist_counts, width=(bin_centers[1]-bin_centers[0])*0.9, alpha=0.6, label="Histogram", color="#6c757d")
    area = np.sum(hist_counts) * (bin_centers[1]-bin_centers[0])
    ax.plot(x_pdf, pdf_theory * area, 'r-', linewidth=2, label="Gaussian Fitting")
    if usl is not None:
        ax.axvline(usl, color='green', linestyle='--', linewidth=1.5, label=f"USL = {usl:.2f}")
    if lsl is not None:
        ax.axvline(lsl, color='orange', linestyle='--', linewidth=1.5, label=f"LSL = {lsl:.2f}")
    stats_text = f"NO.={n_sim}\nAVE={np.mean(results):.2f}\nSTD={np.std(results, ddof=1):.4f}\nMAX={np.max(results):.2f}\nMIN={np.min(results):.2f}"
    ax.text(0.95, 0.95, stats_text, transform=ax.transAxes, fontsize=9, verticalalignment='top', horizontalalignment='right', bbox=dict(boxstyle='round', facecolor='white', alpha=0.8, edgecolor='gray'))
    ax.legend(loc='upper right', bbox_to_anchor=(0.95, 0.72), fontsize=9)
    ax.set_xlabel(output_name, fontsize=11)
    ax.set_ylabel("Frequency", fontsize=11)
    ax.set_title(f"{output_name} Distribution", fontsize=13, fontweight='bold')
    ax.grid(axis='y', linestyle='--', alpha=0.3)
    return fig

def plot_contribution_horizontal(contributions: List[float], param_names: List[str], output_name: str):
    non_zero = [(p, c) for p, c in zip(param_names, contributions) if c > 0]
    if not non_zero:
        fig, ax = plt.subplots(figsize=(8, 4))
        ax.text(0.5, 0.5, "No significant contribution", ha='center', va='center')
        ax.set_title(f"Design Factor Effect % on {output_name}", fontsize=13, fontweight='bold')
        return fig
    names, vals = zip(*non_zero)
    sorted_indices = np.argsort(vals)
    names = [names[i] for i in sorted_indices]
    vals = [vals[i] for i in sorted_indices]
    fig, ax = plt.subplots(figsize=(9, max(4, len(names)*0.4)))
    bars = ax.barh(names, vals, color='#6c757d')
    for bar, val in zip(bars, vals):
        ax.text(val + 0.01, bar.get_y() + bar.get_height()/2, f'{val:.1%}', va='center', fontsize=9)
    ax.set_xlabel("Effect %", fontsize=11)
    ax.set_title(f"Design Factor Effect % on {output_name}", fontsize=13, fontweight='bold')
    ax.set_xlim(0, max(vals) * 1.15)
    ax.grid(axis='x', linestyle='--', alpha=0.5)
    ax.legend().remove()
    return fig

def compute_cpk_ppm(results: np.ndarray, usl: Optional[float], lsl: Optional[float]):
    mean_out = np.mean(results)
    std_out = np.std(results, ddof=1)
    if std_out == 0:
        cpk = None
    else:
        if usl is not None and lsl is not None:
            cpk = min((usl - mean_out) / (3 * std_out), (mean_out - lsl) / (3 * std_out))
        elif usl is not None:
            cpk = (usl - mean_out) / (3 * std_out)
        elif lsl is not None:
            cpk = (mean_out - lsl) / (3 * std_out)
        else:
            cpk = None
    failures_up = np.sum(results > usl) / len(results) * 1e6 if usl is not None else None
    failures_dn = np.sum(results < lsl) / len(results) * 1e6 if lsl is not None else None
    failures_all = None
    if failures_up is not None and failures_dn is not None:
        failures_all = failures_up + failures_dn
    elif failures_up is not None:
        failures_all = failures_up
    elif failures_dn is not None:
        failures_all = failures_dn
    return cpk, failures_all, failures_up, failures_dn

def generate_word_report(raw, usl, lsl, n_sim, seed, formula, params_df, param_letters, analyst_name, analyst_title, output_name):
    results = raw["results"]
    cpk, failures_all, failures_up, failures_dn = compute_cpk_ppm(results, usl, lsl)

    doc = Document()
    title = doc.add_heading(t("report_title").format(output_name), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(t("analyst_info_report"), level=1)
    p = doc.add_paragraph()
    p.add_run(f"{t('analyst_name_report')} {analyst_name if analyst_name else t('not_filled')}").bold = True
    p.add_run(f"\n{t('analyst_title_report')} {analyst_title if analyst_title else t('not_filled')}")
    doc.add_heading(t("sim_settings_report"), level=1)
    doc.add_paragraph(f"{t('output_var')} {output_name}")
    doc.add_paragraph(f"{t('formula_report')} {formula}")
    doc.add_paragraph(f"{t('sim_times')} {n_sim}")
    doc.add_paragraph(f"{t('random_seed_report')} {seed}")
    doc.add_paragraph(f"{t('usl_report')} {usl if usl is not None else t('none')}")
    doc.add_paragraph(f"{t('lsl_report')} {lsl if lsl is not None else t('none')}")
    doc.add_heading(t("param_table"), level=1)
    table = doc.add_table(rows=len(params_df)+1, cols=len(params_df.columns))
    table.style = 'Table Grid'
    for j, col in enumerate(params_df.columns):
        table.cell(0, j).text = col
    for i, row in params_df.iterrows():
        for j, col in enumerate(params_df.columns):
            if col == "分布参数":
                table.cell(i+1, j).text = str(row[col]) if row[col] else "{}"
            else:
                table.cell(i+1, j).text = str(row[col])
    doc.add_heading(t("result_stats").format(output_name), level=1)
    stats_table = doc.add_table(rows=5, cols=2)
    stats_table.style = 'Table Grid'
    stats_table.cell(0, 0).text = t("statistic")
    stats_table.cell(0, 1).text = t("value")
    stats_table.cell(1, 0).text = t("mean_stat")
    stats_table.cell(1, 1).text = f"{raw['mean']:.2f}"
    stats_table.cell(2, 0).text = t("std_stat")
    stats_table.cell(2, 1).text = f"{raw['std']:.4f}"
    stats_table.cell(3, 0).text = t("max_stat")
    stats_table.cell(3, 1).text = f"{raw['max']:.2f}"
    stats_table.cell(4, 0).text = t("min_stat")
    stats_table.cell(4, 1).text = f"{raw['min']:.2f}"
    if cpk is not None:
        stats_table.add_row().cells[0].text = t("cpk_stat")
        stats_table.rows[-1].cells[1].text = f"{cpk:.2f}"
        stats_table.add_row().cells[0].text = t("fail_all")
        stats_table.rows[-1].cells[1].text = f"{failures_all:.2f}" if failures_all is not None else "-"
        stats_table.add_row().cells[0].text = t("fail_up")
        stats_table.rows[-1].cells[1].text = f"{failures_up:.2f}" if failures_up is not None else "-"
        stats_table.add_row().cells[0].text = t("fail_dn")
        stats_table.rows[-1].cells[1].text = f"{failures_dn:.2f}" if failures_dn is not None else "-"
    doc.add_heading(t("histogram_report"), level=1)
    fig_hist = plot_histogram(results, raw["bin_centers"], raw["hist_counts"], raw["x_pdf"], raw["pdf_theory"], usl, lsl, output_name, n_sim)
    buf_hist = BytesIO()
    fig_hist.savefig(buf_hist, format='png', dpi=150, bbox_inches='tight')
    buf_hist.seek(0)
    doc.add_picture(buf_hist, width=Inches(6))
    plt.close(fig_hist)
    doc.add_heading(t("effect_report").format(output_name), level=1)
    fig_barh = plot_contribution_horizontal(raw["contributions"], raw["param_names"], output_name)
    buf_barh = BytesIO()
    fig_barh.savefig(buf_barh, format='png', dpi=150, bbox_inches='tight')
    buf_barh.seek(0)
    doc.add_picture(buf_barh, width=Inches(6))
    plt.close(fig_barh)
    doc.add_heading(t("detail_table"), level=2)
    df_contrib = raw["df_contrib"].copy()
    df_contrib["贡献百分比"] = df_contrib["贡献百分比_显示"]
    contrib_table = doc.add_table(rows=len(df_contrib)+1, cols=2)
    contrib_table.style = 'Table Grid'
    contrib_table.cell(0, 0).text = t("param")
    contrib_table.cell(0, 1).text = t("contribution")
    for i, row in df_contrib.iterrows():
        contrib_table.cell(i+1, 0).text = row["参数"]
        contrib_table.cell(i+1, 1).text = row["贡献百分比"]
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(t("contact_report")).italic = True
    footer.add_run(f"\n{t('report_date').format(datetime.now().strftime('%Y-%m-%d'))}")
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

# ==================== 主函数 ====================
def main():
    handle_payment_callback()
    show_payment_success_dialog()

    st.markdown("""
    <style>
        html, body, .stApp, .stMarkdown, .stText, .stNumberInput, .stSelectbox, .stTextArea, .stDataFrame, .stMetric {
            color: #000000 !important;
        }
        .main-title { font-size: 2.5rem; font-weight: 600; color: #000000; margin-bottom: 1rem; }
        .section-header { font-size: 1.5rem; font-weight: 500; color: #000000; border-left: 5px solid #cccccc; padding-left: 15px; margin: 20px 0 15px 0; }
        .metric-card { background-color: #f8f9fa; border-radius: 10px; padding: 15px; text-align: center; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
        .metric-label { font-size: 1rem; color: #000000; margin-bottom: 5px; }
        .metric-value { font-size: 1.8rem; font-weight: 600; color: #000000; }
        .ppm-table { border-collapse: collapse; width: 100%; margin: 0 auto; }
        .ppm-table th, .ppm-table td { border: 2px solid #000000; padding: 10px 16px; text-align: center; font-size: 1rem; }
        .ppm-table th { background-color: #e9ecef; font-weight: 600; }
        button[data-testid="baseButton-primary"] {
            background-color: #dc3545 !important;
            color: white !important;
            font-weight: 500;
            border-radius: 5px;
            font-size: 1.2rem;
            margin-top: 20px;
            white-space: pre-line !important;
        }
        button[data-testid="baseButton-primary"]:hover { background-color: #c82333 !important; }
        button[data-testid="baseButton-primary"] * { white-space: pre-line !important; }
        .stButton > button:not([data-testid="baseButton-primary"]) {
            background-color: #3498db !important;
            color: white !important;
            font-weight: 500;
            border-radius: 5px;
        }
        .stButton > button:not([data-testid="baseButton-primary"]):hover { background-color: #2980b9 !important; }
        .lang-btn-wrap .stButton button {
            background-color: #dc3545 !important;
            color: white !important;
            font-weight: 500;
            border-radius: 5px;
        }
        .lang-btn-wrap .stButton button:hover { background-color: #c82333 !important; }
        .design-value-card { background-color: #e8f4fd; border-radius: 10px; padding: 15px; margin-top: 15px; text-align: center; border-left: 5px solid #cccccc; }
        .design-value-card strong { font-size: 1.1rem; color: #000000; }
        .design-value-number { font-size: 1.6rem; font-weight: 600; color: #000000; margin-top: 5px; }
        .big-label { font-size: 1.3rem; font-weight: 500; margin-bottom: 5px; color: #000000; }
        .param-letter { font-weight: bold; font-size: 1rem; text-align: center; background-color: #e9ecef; border-radius: 4px; padding: 6px 0; width: 40px; color: #000000; }
        .formula-hint { font-size: 0.9rem; color: #000000; margin-bottom: 5px; }
    </style>
    """, unsafe_allow_html=True)

    # 语言切换 + 设置按钮
    col_lang1, col_lang2, col_lang3, col_gear = st.columns([0.6, 0.15, 0.15, 0.1])
    with col_lang2:
        st.markdown('<div class="lang-btn-wrap">', unsafe_allow_html=True)
        if st.button("中文", key="lang_zh", use_container_width=True):
            st.session_state.lang = "zh"
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    with col_lang3:
        st.markdown('<div class="lang-btn-wrap">', unsafe_allow_html=True)
        if st.button("English", key="lang_en", use_container_width=True):
            st.session_state.lang = "en"
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    with col_gear:
        if st.button("⚙️", key="settings_btn", use_container_width=True):
            admin_settings_dialog()

    st.markdown(f'<div class="main-title">{t("title")}</div>', unsafe_allow_html=True)
    st.markdown(t("subtitle"))

    # 侧边栏
    with st.sidebar:
        st.markdown(f"## {t('sim_settings')}")
        n_sim = st.number_input(t("trail_number"), min_value=100, max_value=100000, value=1000, step=100)
        st.markdown(f"#### {t('spec_limits')}")
        usl_sidebar = st.text_input(t("usl"), value=st.session_state.usl_str, key="usl_sidebar", on_change=sync_usl_from_sidebar)
        lsl_sidebar = st.text_input(t("lsl"), value=st.session_state.lsl_str, key="lsl_sidebar", on_change=sync_lsl_from_sidebar)
        st.session_state.usl_str = usl_sidebar
        st.session_state.lsl_str = lsl_sidebar
        seed = st.number_input(t("random_seed"), value=38, step=1)

        st.markdown("---")
        st.markdown(f"### {t('about_system')}")
        st.markdown(t("about_desc1"))
        st.markdown(t("about_desc2"))
        st.markdown(t("output_title"))
        st.markdown(t("output1"))
        st.markdown(t("output2"))
        st.markdown(t("output3"))

        st.markdown("---")
        st.markdown(f"### {t('analyst_info')}")
        analyst_name = st.text_input(t("analyst_name"), value=st.session_state.analyst_name, key="analyst_name_input")
        analyst_title = st.text_input(t("analyst_title"), value=st.session_state.analyst_title, key="analyst_title_input")
        st.session_state.analyst_name = analyst_name
        st.session_state.analyst_title = analyst_title

        st.markdown("---")
        st.markdown(f"### 🔑 {t('license_info')}")
        new_report_key = st.text_input(t("report_key_label"), value=st.session_state.current_report_key, type="password", key="report_key_input", placeholder="输入授权码后按 Enter")
        if new_report_key != st.session_state.current_report_key:
            st.session_state.current_report_key = new_report_key
            if new_report_key:
                valid, remaining, expiry_str, _ = activate_license(new_report_key)
                if valid:
                    st.success(f"授权成功！剩余 {remaining} 次，有效期至 {expiry_str[:10]}" if st.session_state.lang=="zh" else f"Success! {remaining} uses left, valid until {expiry_str[:10]}")
                    st.rerun()
                else:
                    st.error("授权码无效或已过期" if st.session_state.lang=="zh" else "Invalid or expired license key")
                    st.session_state.current_report_key = ""
                    st.rerun()
            else:
                st.rerun()
        remaining_str, expiry_str = get_remaining_info(st.session_state.current_report_key)
        st.write(f"{t('remaining_label')}: {remaining_str}")
        if expiry_str not in ("试用剩余次数", "Trial left"):
            st.write(f"{t('expiry_label')}: {expiry_str}")
        if not is_premium_user(st.session_state.current_report_key):
            st.warning(t("trial_warning").format(st.session_state.trial_uses_left))

        if st.button(t("purchase_button"), use_container_width=True):
            purchase_dialog()

        st.markdown("---")
        st.markdown(f"**{t('contact')}**")
        st.markdown(t("email"))

    # ==================== 参数输入表格 ====================
    st.markdown(f'<div class="section-header">{t("param_input")}</div>', unsafe_allow_html=True)
    header_cols = st.columns([0.3, 1.5, 1, 1, 1.2, 0.3])
    header_cols[0].markdown(f"**{t('letter')}**")
    header_cols[1].markdown(f"**{t('param_name')}**")
    header_cols[2].markdown(f"**{t('mean')}**")
    header_cols[3].markdown(f"**{t('std')}**")
    header_cols[4].markdown(f"**{t('distribution')}**")
    header_cols[5].markdown(f"**{t('delete')}**")

    rows_data = []
    distributions_list = get_distributions()
    for idx, row in st.session_state.params.iterrows():
        letter = chr(ord('A') + idx)
        cols = st.columns([0.3, 1.5, 1, 1, 1.2, 0.3])
        with cols[0]:
            st.markdown(f'<div class="param-letter">{letter}</div>', unsafe_allow_html=True)
        with cols[1]:
            name = st.text_input("", value=row["参数名称"], key=f"param_name_{idx}", label_visibility="collapsed")
        with cols[2]:
            mean_val = st.number_input("", value=float(row["均值(Typ)"]), step=1.0, key=f"param_mean_{idx}", label_visibility="collapsed")
        with cols[3]:
            std_val = st.number_input("", value=float(row["标准差(Std)"]), step=0.01, format="%.4f", key=f"param_std_{idx}", label_visibility="collapsed")
        with cols[4]:
            dist_val = st.selectbox("", distributions_list, index=distributions_list.index(row["分布"]) if row["分布"] in distributions_list else 0, key=f"param_dist_{idx}", label_visibility="collapsed")
        with cols[5]:
            delete = st.button("🗑️", key=f"del_{idx}")

        current_dist_params = row.get("分布参数", {}) if isinstance(row.get("分布参数"), dict) else {}
        if dist_val in [t("dist_uniform"), t("dist_lognorm"), t("dist_weibull"), t("dist_tri")]:
            if dist_val == t("dist_uniform") and "low" not in current_dist_params:
                current_dist_params["low"] = mean_val - 3 * std_val
                current_dist_params["high"] = mean_val + 3 * std_val
            elif dist_val == t("dist_lognorm") and "mean_log" not in current_dist_params:
                current_dist_params["mean_log"] = 0.0
                current_dist_params["sigma_log"] = 1.0
            elif dist_val == t("dist_weibull") and "shape" not in current_dist_params:
                current_dist_params["shape"] = 1.0
                current_dist_params["scale"] = 1.0
            elif dist_val == t("dist_tri") and "left" not in current_dist_params:
                current_dist_params["left"] = mean_val - 3 * std_val
                current_dist_params["mode"] = mean_val
                current_dist_params["right"] = mean_val + 3 * std_val

        need_expand = dist_val in [t("dist_uniform"), t("dist_lognorm"), t("dist_weibull"), t("dist_tri")]
        if need_expand:
            with st.expander(t("configure").format(dist_val), expanded=True):
                if dist_val == t("dist_uniform"):
                    low = st.number_input(t("uniform_low"), value=float(current_dist_params.get("low", mean_val - 3*std_val)), key=f"uniform_low_{idx}", step=0.1)
                    high = st.number_input(t("uniform_high"), value=float(current_dist_params.get("high", mean_val + 3*std_val)), key=f"uniform_high_{idx}", step=0.1)
                    if low >= high:
                        st.error(t("error_low_high"))
                    else:
                        current_dist_params["low"] = low
                        current_dist_params["high"] = high
                elif dist_val == t("dist_lognorm"):
                    mean_log = st.number_input(t("lognorm_meanlog"), value=float(current_dist_params.get("mean_log", 0.0)), key=f"lognorm_meanlog_{idx}", step=0.1)
                    sigma_log = st.number_input(t("lognorm_sigmalog"), value=float(current_dist_params.get("sigma_log", 1.0)), key=f"lognorm_sigmalog_{idx}", step=0.05, format="%.3f")
                    if sigma_log <= 0:
                        st.error(t("error_sigma"))
                    else:
                        current_dist_params["mean_log"] = mean_log
                        current_dist_params["sigma_log"] = sigma_log
                elif dist_val == t("dist_weibull"):
                    shape = st.number_input(t("weibull_shape"), value=float(current_dist_params.get("shape", 1.0)), key=f"weibull_shape_{idx}", step=0.1, min_value=0.1)
                    scale = st.number_input(t("weibull_scale"), value=float(current_dist_params.get("scale", 1.0)), key=f"weibull_scale_{idx}", step=0.1, min_value=0.1)
                    if shape <= 0 or scale <= 0:
                        st.error(t("error_weibull"))
                    else:
                        current_dist_params["shape"] = shape
                        current_dist_params["scale"] = scale
                elif dist_val == t("dist_tri"):
                    left = st.number_input(t("tri_left"), value=float(current_dist_params.get("left", mean_val - 3*std_val)), key=f"tri_left_{idx}", step=0.1)
                    mode = st.number_input(t("tri_mode"), value=float(current_dist_params.get("mode", mean_val)), key=f"tri_mode_{idx}", step=0.1)
                    right = st.number_input(t("tri_right"), value=float(current_dist_params.get("right", mean_val + 3*std_val)), key=f"tri_right_{idx}", step=0.1)
                    if not (left <= mode <= right):
                        st.error(t("error_tri"))
                    else:
                        current_dist_params["left"] = left
                        current_dist_params["mode"] = mode
                        current_dist_params["right"] = right

                fig, ax = plt.subplots(figsize=(4, 2))
                plot_pdf(dist_val, mean_val, std_val, current_dist_params, ax)
                st.pyplot(fig)
                plt.close(fig)

        rows_data.append((name, mean_val, std_val, dist_val, current_dist_params, delete, letter))

    new_params = []
    for (name, mean_val, std_val, dist_val, dist_params, delete, letter) in rows_data:
        if not delete:
            new_params.append({
                "参数名称": name,
                "均值(Typ)": mean_val,
                "标准差(Std)": std_val,
                "分布": dist_val,
                "分布参数": dist_params
            })
    if st.button(t("add_row"), use_container_width=True):
        new_params.append({
            "参数名称": t("new_param_default"),
            "均值(Typ)": 0.0,
            "标准差(Std)": 0.0,
            "分布": t("dist_full"),
            "分布参数": {}
        })

    st.session_state.params = pd.DataFrame(new_params)
    update_param_letters()

    # ==================== 公式定义区域 ====================
    st.markdown(f'<div class="section-header">{t("formula_def")}</div>', unsafe_allow_html=True)
    st.markdown(f'<span class="big-label">{t("design_var_name")}</span>', unsafe_allow_html=True)
    output_name = st.text_input("", value=st.session_state.output_name, key="output_name_input", label_visibility="collapsed")
    st.session_state.output_name = output_name if output_name.strip() else "Output"
    st.markdown(f'<span class="big-label">{t("formula_label")}</span>', unsafe_allow_html=True)
    st.markdown(f'<div class="formula-hint">{t("formula_hint")}</div>', unsafe_allow_html=True)
    formula = st.text_area("", value=st.session_state.formula, height=100, key="formula_input", label_visibility="collapsed")
    st.session_state.formula = formula
    st.caption(t("formula_supported"))

    design_val = compute_design_value(st.session_state.params, formula, st.session_state.param_letters)
    if design_val is not None and not np.isnan(design_val):
        st.markdown(f"""
        <div class="design-value-card">
            <strong>{t("design_value")}</strong><br>
            <span class="design-value-number">{output_name} = {design_val:.2f}</span>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.warning(t("formula_invalid"))

    # ==================== 模拟按钮 ====================
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button(t("start_sim"), type="primary", use_container_width=True):
            if not is_premium_user(st.session_state.current_report_key):
                if st.session_state.trial_uses_left <= 0:
                    st.error(t("analyze_disabled"))
                    purchase_dialog()
                    st.stop()
            if not consume_usage(st.session_state.current_report_key):
                st.error(t("analyze_disabled"))
                purchase_dialog()
                st.stop()

            if st.session_state.params.isnull().values.any():
                st.error(t("formula_invalid"))
                st.stop()
            param_names = st.session_state.params["参数名称"].astype(str).tolist()
            if len(set(param_names)) != len(param_names):
                st.error(t("formula_invalid"))
                st.stop()
            if not formula.strip():
                st.error(t("formula_invalid"))
                st.stop()

            with st.spinner(t("start_sim")):
                sim_res = run_monte_carlo(st.session_state.params, formula, n_sim, st.session_state.param_letters, seed)
            if sim_res is None:
                st.stop()

            with st.spinner(t("start_sim")):
                df_contrib, contributions, param_names = sensitivity_analysis(st.session_state.params, formula, n_sim, st.session_state.param_letters, seed)

            st.session_state.sim_results_raw = {
                "results": sim_res["results"],
                "samples": sim_res["samples"],
                "mean": sim_res["mean"],
                "std": sim_res["std"],
                "max": sim_res["max"],
                "min": sim_res["min"],
                "hist_counts": sim_res["hist_counts"],
                "bin_edges": sim_res["bin_edges"],
                "bin_centers": sim_res["bin_centers"],
                "x_pdf": sim_res["x_pdf"],
                "pdf_theory": sim_res["pdf_theory"],
                "param_names": sim_res["param_names"],
                "df_contrib": df_contrib,
                "contributions": contributions,
                "params_df": st.session_state.params,
                "output_name": output_name,
                "formula": formula,
            }

    # ==================== 结果显示 ====================
    if st.session_state.sim_results_raw is not None:
        raw = st.session_state.sim_results_raw
        results = raw["results"]
        output_name = raw["output_name"]
        usl = parse_limit(st.session_state.usl_str)
        lsl = parse_limit(st.session_state.lsl_str)
        cpk, failures_all, failures_up, failures_dn = compute_cpk_ppm(results, usl, lsl)

        st.markdown(f'<div class="section-header">{t("sim_result").format(output_name)}</div>', unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(f'<div class="metric-card"><div class="metric-label">{t("mean_val").format(output_name)}</div><div class="metric-value">{raw["mean"]:.2f}</div></div>', unsafe_allow_html=True)
            st.markdown(f'<div class="metric-card"><div class="metric-label">{t("std_val").format(output_name)}</div><div class="metric-value">{raw["std"]:.4f}</div></div>', unsafe_allow_html=True)
        with col2:
            st.markdown(f'<div class="metric-card"><div class="metric-label">{t("max_val")}</div><div class="metric-value">{raw["max"]:.2f}</div></div>', unsafe_allow_html=True)
            st.markdown(f'<div class="metric-card"><div class="metric-label">{t("min_val")}</div><div class="metric-value">{raw["min"]:.2f}</div></div>', unsafe_allow_html=True)
        with col3:
            if cpk is not None:
                st.markdown(f'<div class="metric-card"><div class="metric-label">{t("cpk_val")}</div><div class="metric-value">{cpk:.2f}</div></div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="metric-card"><div class="metric-label">{t("cpk_val")}</div><div class="metric-value">-</div></div>', unsafe_allow_html=True)

        st.markdown(f'<div class="section-header">{t("failure_ppm")}</div>', unsafe_allow_html=True)
        st.caption(t("ppm_hint"))
        col_left, col_right = st.columns([1, 2])
        with col_left:
            main_usl = st.text_input(t("usl"), value=st.session_state.usl_str, key="main_usl", on_change=sync_usl_from_main)
            main_lsl = st.text_input(t("lsl"), value=st.session_state.lsl_str, key="main_lsl", on_change=sync_lsl_from_main)
            st.session_state.usl_str = main_usl
            st.session_state.lsl_str = main_lsl
            usl = parse_limit(main_usl)
            lsl = parse_limit(main_lsl)
            cpk, failures_all, failures_up, failures_dn = compute_cpk_ppm(results, usl, lsl)
        with col_right:
            if cpk is not None:
                def fmt(v): return f"{v:.2f}" if v is not None else "-"
                st.markdown(f"""
                <table class="ppm-table">
                    <tr><th>CPK</th><th>Failure All</th><th>Failure Up</th><th>Failure Dn</th></tr>
                    <tr><td style="text-align:center">{fmt(cpk)}</td><td style="text-align:center">{fmt(failures_all)}</td><td style="text-align:center">{fmt(failures_up)}</td><td style="text-align:center">{fmt(failures_dn)}</td></tr>
                </table>
                """, unsafe_allow_html=True)
            else:
                st.info(t("no_limits"))

        st.markdown(f"### {t('histogram').format(output_name)}")
        fig_hist = plot_histogram(results, raw["bin_centers"], raw["hist_counts"], raw["x_pdf"], raw["pdf_theory"], usl, lsl, output_name, n_sim)
        st.pyplot(fig_hist)
        st.caption(t("hist_caption").format(output_name))

        st.markdown(f"### {t('effect_chart').format(output_name)}")
        fig_barh = plot_contribution_horizontal(raw["contributions"], raw["param_names"], output_name)
        st.pyplot(fig_barh)
        st.caption(t("effect_caption"))

        with st.expander(t("view_contrib")):
            st.dataframe(raw["df_contrib"][["参数", "贡献百分比_显示"]].rename(columns={"贡献百分比_显示": "贡献百分比"}), use_container_width=True)

        with st.expander(t("view_data")):
            samples_df = pd.DataFrame(raw["samples"], columns=raw["param_names"])
            samples_df[output_name] = results
            st.dataframe(samples_df.round(2), use_container_width=True, height=400)
            csv = samples_df.to_csv(index=False, float_format="%.6f")
            st.download_button(t("download_csv"), data=csv, file_name=f"monte_carlo_data_{output_name}.csv", mime="text/csv")

        if st.button(t("download_report")):
            if not is_premium_user(st.session_state.current_report_key):
                st.error(t("need_license"))
                purchase_dialog()
            else:
                doc_bytes = generate_word_report(raw, usl, lsl, n_sim, seed, formula, st.session_state.params, st.session_state.param_letters, st.session_state.analyst_name, st.session_state.analyst_title, output_name)
                date_str = datetime.now().strftime("%Y%m%d")
                st.download_button(t("download_report"), data=doc_bytes, file_name=f"DFSS_Report_{output_name}_{date_str}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.success(t("success"))

if __name__ == "__main__":
    main()
